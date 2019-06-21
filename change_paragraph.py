import docx

class Word(object):

    def input_data(paragraph,data,value):

        position=paragraph.text.find(data)+len(data)
        length=len(value)
        new = []
        for s in value:
            new.append(s)
        position_at=position
        for letter in new:        
            inline =paragraph.runs
            begin=0
            for i in range(len(inline)):
                if begin<=position_at<begin+len(inline[i].text) :
                    offset=position_at-begin
                    new = []
                    for s in inline[i].text:
                        new.append(s)
                    new[offset]=letter
                    inline[i].text=''.join(new)
                    position_at=position_at+1;
                 
                    begin_align=0
                    for j in range(len(inline)):
                        if begin_align<=position_at<begin_align+len(inline[j].text) :
                            offset_align=position_at-begin_align
                            if letter.isupper() | letter.isdigit() :
                                
                                inline[j].text=inline[j].text[:offset_align]+inline[j].text[offset_align+1:]
                            break    
                        begin_align=begin_align+len(inline[j].text)    
                       

                    break 
                begin=begin+len(inline[i].text)
        return paragraph

    def input_table_data(doc,data,value,maximum):  

        count=1 
        value_fill=False 
        for table in doc.tables:  # 遍历所有表格
            if count>maximum:            
                break
            count=count+1  
          
            print ('----table------')
            for row in table.rows:  # 遍历表格的所有行
                # row_str = '\t'.join([cell.text for cell in row.cells])  # 一行数据
                # print row_str
                for cell in row.cells:
                    print(cell.text, '\t')
                    if value_fill==True:                   
                       cell.text=value
                       return doc                      

                    if cell.text==data: 
                          value_fill=True    
        return doc    

if __name__ == "__main__":

    doc = docx.Document('Field Problem Feedback-Template.docx')
    word=Word
    word.input_table_data(doc,"City","beijing",4)


'''
    doc = docx.Document('Field Problem Feedback-Template.docx')

    for j in  doc.paragraphs:

        num=0
        for i in range(len(j.runs)):
       
            num=len(j.runs[i].text)+num
            print(j.runs[i].text)
            print("size=",len(j.runs[i].text))
        
        print("number=",num)
        #break

    word=Word

    data="Product line ："
    value="SCHINDLER 3300 "

    doc.paragraphs[0]=word.input_data(doc.paragraphs[1],data,value)
    doc.save('tmp.docx')

    doc = docx.Document('tmp.docx')

    for j in  doc.paragraphs:

        num=0
        for i in range(len(j.runs)):
       
            num=len(j.runs[i].text)+num
            print(j.runs[i].text)
            print("size=",len(j.runs[i].text))
        
        print("number=",num)
        #break

'''




















