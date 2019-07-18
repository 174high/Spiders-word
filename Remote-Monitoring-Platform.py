from PyQt5.QtCore import pyqtSlot, QIODevice, QByteArray 
from PyQt5.QtSerialPort import QSerialPortInfo, QSerialPort
from PyQt5.QtWidgets import QWidget, QMessageBox,QProgressDialog
from PyQt5 import QtGui
from PyQt5.QtCore import *
from PyQt5.QtWebEngineWidgets import *
from PyQt5.QtCore import QTimer

import pychrome
from bs4 import BeautifulSoup
import requests
import re

import datetime
import docx
import signal
import os
import subprocess
import _thread
import shutil

from change_paragraph import Word 
from Ui_FormMonitor import Ui_Form 
from run_chrome import stop_chrome,run_chrome
#from excel import merge_file
from PythonHeadlessChrome.download import RMPDownload

class Window(QWidget, Ui_Form):

    def __init__(self, *args, **kwargs):
        super(Window, self).__init__(*args, **kwargs)
        self.setupUi(self)

        self.browser = QWebEngineView(self)
        url="http://localhost:8080/"
       
        self.equipments={"54001574","54001575","54001576","54001577","54001580","54001581","54001582","54001583","54001586","54001587","54001660","54001661","54001662","54001663","54001664","54001665","54001666","54001667","54001668","54001675"}
#        url="https://www.baidu.com"
        self.browser.load(QUrl(url))
#        self.browser.load(QUrl.fromLocalFile(
#        os.path.abspath('data/javascript.html')))        
        
        stop_chrome()
        run_chrome()

        run_num=0
        result=self.start_chrome()
        while(result):
            print("starting ......")
            stop_chrome()
            run_chrome()
            result=self.start_chrome() 
            if run_num>5: 
                exit(0)

        self.gridLayout_5.addWidget(self.browser)
        self.pushButton_2.clicked.connect(self.watch_list)
        self.pushButton.clicked.connect(self.process_one_equipment)
        self.process_log.clicked.connect(self.generate_log)
        self.process_event.clicked.connect(self.generate_event)
        self.event_summary.clicked.connect(self.genarate_event_summary)
        self.log_summary.clicked.connect(self.genarate_log_summary)
        self.download_excel.clicked.connect(self.rmp_download_excel)

    def start_chrome(self):

        try:
        # create a browser instance
            self.browser_chrome = pychrome.Browser(url="http://127.0.0.1:9235")
        # create a tab'
            self.tab = self.browser_chrome.new_tab()
	# start the tab 
            self.tab.start()
        # call method
            self.tab.Network.enable() 
        except: 
            print("run pychrome.Browser error")
            return 1

        return 0       

    def call_web(self,equipment,progress_in,progress_out):
        for i in range(progress_in,progress_out):            
            self.progress.setValue(i) 
            if self.progress.wasCanceled():
                QMessageBox.warning(self,"提示","操作失败") 
                return 2

        try:
            # call method with timeout
            self.tab.call_method('Page.navigate',url="http://rmp.global.schindler.com/Equipment/EquipmentMain/EquipmentDetails/?sapSys=ZAP&equnr="+equipment, _timeout=2)
	    # wait for loading
            self.tab.wait(1)

        except:
            #stop_chrome()
            #run_chrome()
            print ("Error: run chrome error!")
            return 1

        try:
            html = self.tab.Runtime.evaluate(expression="document.documentElement.outerHTML")
            soup = BeautifulSoup(((html['result'])['value']),"html.parser")
            print(soup.prettify())

            #print(soup)
            #print(soup.select("#ProductLineDesc"))
            #print(type((soup.select("#ProductLineDesc"))[0]['value']))
            #print(soup.title)	

            doc = docx.Document('./device-file/'+'Field Problem Feedback-Template.docx')

            data="Originator: "
            doc.paragraphs[0]=Word.input_data(doc.paragraphs[0],data,"JOHNNY")
            data="Entry Date："
            doc.paragraphs[0]=Word.input_data(doc.paragraphs[0],data,datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S'))
            data="Phone/Mail:"
            doc.paragraphs[0]=Word.input_data(doc.paragraphs[0],data,(soup.select("#TaPhoneNumber"))[0]['value'])
            data="Product line ："        
            doc.paragraphs[1]=Word.input_data(doc.paragraphs[1],data,(soup.select("#ProductLineDesc"))[0]['value'])
            data="Zone: "
            doc.paragraphs[1]=Word.input_data(doc.paragraphs[1],data,"N/A")
            data="Clams / PROCITS / PROMs Nr："
            doc.paragraphs[1]=Word.input_data(doc.paragraphs[1],data,"N/A")
            data="Commission No："
            doc.paragraphs[2]=Word.input_data(doc.paragraphs[2],data,(soup.select("#CommisionNumber"))[0]['value'])
            data="Handover to EI: "
            doc.paragraphs[2]=Word.input_data(doc.paragraphs[2],data,"N/A")
            data="Maintenance by Schindler："
            doc.paragraphs[2]=Word.input_data(doc.paragraphs[2],data,"N/A")
            data="Affected Component："
            doc.paragraphs[3]=Word.input_data(doc.paragraphs[3],data,"N/A") 
            data="Number of Units："
            doc.paragraphs[3]=Word.input_data(doc.paragraphs[3],data,"N/A") 
            data="Other product line affected: "
            doc.paragraphs[3]=Word.input_data(doc.paragraphs[3],data,"N/A") 
            doc=Word.input_table_data(doc,"City","beijing",4)

#           doc=Word.input_table_data(doc,"Machine Type",(soup.select("#Machine type"))[0]['value'],4)

            doc.save("./result/"+'Feedback-'+equipment+'.docx')

            doc = docx.Document('./device-file/'+'rmp info-Template.docx')
            doc=Word.input_table_data(doc,"Equipment #",(soup.select("#Equnr"))[0]['value'],4) 
            doc=Word.input_table_data(doc,"Product Line",(soup.select("#ProductLineDesc"))[0]['value'],4)
            doc=Word.input_table_data(doc,"Address",(soup.select("#Address"))[0]['value'],4)
            doc=Word.input_table_data(doc,"City",(soup.select("#CityAndZip"))[0]['value'],4)
            doc=Word.input_table_data(doc,"CompanyCode",(soup.select("#BukrsDesc"))[0]['value'],4)
            doc=Word.input_table_data(doc,"Controller Type",(soup.select("#ControllerTypeSwVersion"))[0]['value'],4)
            doc=Word.input_table_data(doc,"Gateway Type",(soup.select("#GatewayType"))[0]['value'],4)
            doc=Word.input_table_data(doc,"Serial Number",(soup.select("#SerialNumber"))[0]['value'],4)
            doc=Word.input_table_data(doc,"TM SW Version",(soup.select("#TmSwVersion"))[0]['value'],4)
            doc=Word.input_table_data(doc,"TM State",(soup.select("#TmState"))[0]['value'],4)
            doc=Word.input_table_data(doc,"Environment",(soup.select("#Environment"))[0]['value'],4)
            doc=Word.input_table_data(doc,"Drive Type",(soup.select("#DriveType"))[0]['value'],4)
            doc=Word.input_table_data(doc,"Door Type",(soup.select("#DoorType"))[0]['value'],4)
            doc=Word.input_table_data(doc,"Machine Type",(soup.select("#MachineType"))[0]['value'],4)
            doc=Word.input_table_data(doc,"Workcenter",(soup.select("#Workcenter"))[0]['value'],4)
            doc=Word.input_table_data(doc,"Branch",(soup.select("#Branch"))[0]['value'],4)
            doc=Word.input_table_data(doc,"Commission #",(soup.select("#CommisionNumber"))[0]['value'],4)
            doc=Word.input_table_data(doc,"Equipment Title",(soup.select("#Description"))[0]['value'],4)

            data="Originator: "
            doc.paragraphs[0]=Word.input_data(doc.paragraphs[0],data,"JOHNNY")
            data="Entry Date："
            doc.paragraphs[0]=Word.input_data(doc.paragraphs[0],data,datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S'))
            data="Phone/Mail:"
            doc.paragraphs[0]=Word.input_data(doc.paragraphs[0],data,(soup.select("#TaPhoneNumber"))[0]['value'])
            doc.save("./result/"+"rmp info-result-"+equipment+".docx")

            print("address=",(soup.select("#CityAndZip"))[0]['value'],(soup.select("#Address"))[0]['value'])
        except:
            #stop_chrome()
            #run_chrome()
            print ("Error: process html error!")
            print ("issue solved 1:  rmp info-Template.docx,Field Problem Feedback-Template.docx needed !!! " )
            return 1

        # stop the tab (stop handle events and stop recv message from chrome)
        #self.tab.stop()
        #self.browser_chrome.close_tab(self.tab)

        return 0

#    def request_will_be_sent(**kwargs):
#        print("loading: %s" % kwargs.get('request').get('url'))

    def watch_list(self):
        print("watch list")
  
        print(equipments)

        num = 10000
        self.progress = QProgressDialog(self)
        self.progress.setWindowTitle("请稍等")  
        self.progress.setLabelText("正在操作...")
        self.progress.setCancelButtonText("取消")
        self.progress.setMinimumDuration(5)
        self.progress.setWindowModality(Qt.WindowModal)
        self.progress.setRange(0,num) 
        self.progress.setValue(0)  

        for i in range(0,1000):            
            self.progress.setValue(i) 
            if self.progress.wasCanceled():
                QMessageBox.warning(self,"提示","操作失败") 
                return 0

        progress_bar=1000
                     
        for equipment in self.equipments:
        
            result=self.call_web(equipment,progress_bar,progress_bar+400)
            print("result:",result) 
            if result == 2:
                return 0
            progress_bar=progress_bar+400
            run_num=0
            while result  :
                result=self.call_web(equipment,progress_bar,progress_bar)
                print("result:",result) 
                if result == 2:
                    return 0
                run_num=run_num+1
                if run_num > 10:
                    QMessageBox.warning(self,"提示","操作失败,请重启软件") 
                    return 0

            if self.progress.wasCanceled():
                QMessageBox.warning(self,"提示","操作失败") 
                return 0

        for i in range(9000,num):            
            self.progress.setValue(i) 
            if self.progress.wasCanceled():
                QMessageBox.warning(self,"提示","操作失败") 
                return 0

        self.progress.setValue(num)
        QMessageBox.information(self,"提示","操作成功")

    def process_one_equipment(self):
        equipment = self.textEdit.toPlainText()
        print(equipment)

        num = 10000
        self.progress = QProgressDialog(self)
        self.progress.setWindowTitle("请稍等")  
        self.progress.setLabelText("正在操作...")
        self.progress.setCancelButtonText("取消")
        self.progress.setMinimumDuration(5)
        self.progress.setWindowModality(Qt.WindowModal)
        self.progress.setRange(0,num) 
        self.progress.setValue(0)  

        for i in range(0,3000):            
            self.progress.setValue(i) 
            if self.progress.wasCanceled():
                QMessageBox.warning(self,"提示","操作失败") 
                return                 

        result=self.call_web(equipment,3000,4000)
        print("result:",result) 
        if result == 2:
            return 0
        run_num=0
        while result  :
            result=self.call_web(equipment,4000,4000)
            print("result:",result) 
            if result == 2:
                return 
            run_num=run_num+1
            if run_num > 10:
               QMessageBox.warning(self,"提示","操作失败,请重启软件") 
               return 

        for i in range(4000,num):            
            self.progress.setValue(i) 
            if self.progress.wasCanceled():
                QMessageBox.warning(self,"提示","操作失败") 
                return 

        self.progress.setValue(num)
        QMessageBox.information(self,"提示","操作成功")
         
    def run_excel(self,cmd, delay):
        subprocess.call("excel"+" "+cmd, shell=True)
        print("test")

    def generate_log(self):
        print("generate_log")
#        merge_file("./device-log/","merge-log.xlsx")
        try:
            _thread.start_new_thread(self.run_excel,("log", 4, ))
        except:
            print ("Error: can't run thread")   
        print("continue ....")
        self.progress_1 = QProgressDialog(self)
        self.progress_1.setWindowTitle("请稍等")  
        self.progress_1.setLabelText("正在操作...")
        self.progress_1.setCancelButtonText("取消")
        self.progress_1.setMinimumDuration(5)
        self.progress_1.setWindowModality(Qt.WindowModal)
        self.progress_1.setRange(0,10000) 
        self.progress_1.setValue(0)  

        for i in range(0,10000):            
            self.progress_1.setValue(i) 
            if self.progress_1.wasCanceled():
                QMessageBox.warning(self,"提示","操作失败") 
                return   
         
        self.progress_1.setValue(10000)
        QMessageBox.information(self,"提示","程序正在后台运行,请耐心等待")

    def generate_event(self):
        print("generate_event")
#        merge_file("./device-event/","merge-event.xlsx")
        try:
            _thread.start_new_thread(self.run_excel,("event", 4, ))
        except:
            print ("Error: can't run thread")   
        print("continue ....")
        self.progress_1 = QProgressDialog(self)
        self.progress_1.setWindowTitle("请稍等")  
        self.progress_1.setLabelText("正在操作...")
        self.progress_1.setCancelButtonText("取消")
        self.progress_1.setMinimumDuration(5)
        self.progress_1.setWindowModality(Qt.WindowModal)
        self.progress_1.setRange(0,10000) 
        self.progress_1.setValue(0)  

        for i in range(0,10000):            
            self.progress_1.setValue(i) 
            if self.progress_1.wasCanceled():
                QMessageBox.warning(self,"提示","操作失败") 
                return   
         
        self.progress_1.setValue(10000)
        QMessageBox.information(self,"提示","程序正在后台运行,请耐心等待")

    def genarate_log_summary(self):

        print("genarate log summary")
        try:
            _thread.start_new_thread(self.run_excel,("summary-log", 4, ))
        except:
            print ("Error: can't run thread")   
        print("continue ....")
        self.progress_1 = QProgressDialog(self)
        self.progress_1.setWindowTitle("请稍等")  
        self.progress_1.setLabelText("正在操作...")
        self.progress_1.setCancelButtonText("取消")
        self.progress_1.setMinimumDuration(5)
        self.progress_1.setWindowModality(Qt.WindowModal)
        self.progress_1.setRange(0,10000) 
        self.progress_1.setValue(0)  

        for i in range(0,10000):            
            self.progress_1.setValue(i) 
            if self.progress_1.wasCanceled():
                QMessageBox.warning(self,"提示","操作失败") 
                return   
         
        self.progress_1.setValue(10000)
        QMessageBox.information(self,"提示","程序正在后台运行,请耐心等待")

    def genarate_event_summary(self):
        print("genarate event summary")
        try:
            _thread.start_new_thread(self.run_excel,("summary-event", 4, ))
        except:
            print ("Error: can't run thread")   
        print("continue ....")
        self.progress_1 = QProgressDialog(self)
        self.progress_1.setWindowTitle("请稍等")  
        self.progress_1.setLabelText("正在操作...")
        self.progress_1.setCancelButtonText("取消")
        self.progress_1.setMinimumDuration(5)
        self.progress_1.setWindowModality(Qt.WindowModal)
        self.progress_1.setRange(0,10000) 
        self.progress_1.setValue(0)  

        for i in range(0,10000):            
            self.progress_1.setValue(i) 
            if self.progress_1.wasCanceled():
                QMessageBox.warning(self,"提示","操作失败") 
                return   
         
        self.progress_1.setValue(10000)
        QMessageBox.information(self,"提示","程序正在后台运行,请耐心等待")

    def run_download(self,cmd, delay):
    #    self.equipments
        for equipment in self.equipments:
            dirname, filename = os.path.split(os.path.abspath(__file__)) 
            print("dir=",dirname)
            t=RMPDownload(dirname+"/")
            t.download_by_quip(equipment)

            fileList=os.listdir(dirname)
            for file in fileList:
               print("file name=",file)
               if file.find("EventsExport") >=0 :
                   shutil.move(file,"./device-event/")
               elif file.find("MessagesExport") >=0 :
                   shutil.move(file,"./device-log/")

    def rmp_download_excel(self):

        try:
            _thread.start_new_thread(self.run_download,("log", 4, ))
        except:
            print ("Error: can't run thread")   
        print("continue ....")
        self.progress_1 = QProgressDialog(self)
        self.progress_1.setWindowTitle("请稍等")  
        self.progress_1.setLabelText("正在操作...")
        self.progress_1.setCancelButtonText("取消")
        self.progress_1.setMinimumDuration(5)
        self.progress_1.setWindowModality(Qt.WindowModal)
        self.progress_1.setRange(0,10000) 
        self.progress_1.setValue(0)  

        for i in range(0,10000):            
            self.progress_1.setValue(i) 
            if self.progress_1.wasCanceled():
                QMessageBox.warning(self,"提示","操作失败") 
                return   
         
        self.progress_1.setValue(10000)
        QMessageBox.information(self,"提示","程序正在后台运行,请耐心等待")

def signal_handler(signal,frame):
    print('You pressed Ctrl+C!')
#    tab.stop()
    sys.exit(0)


if __name__ == '__main__':
    import sys
    import cgitb
    sys.excepthook = cgitb.enable(1, None, 5, '')
    from PyQt5.QtWidgets import QApplication

    app = QApplication(sys.argv)
    w = Window()
    w.setFixedSize(850,600)
    signal.signal(signal.SIGINT,signal_handler)
    w.show()
    sys.exit(app.exec_())












